import hashlib
import io
import re
from dataclasses import dataclass, field
from typing import Iterable


@dataclass
class RedactionOptions:
    redact_emails: bool = True
    redact_phones: bool = True
    redact_ssn_tax_ids: bool = True
    redact_addresses: bool = True
    redact_names: bool = True
    redact_dob: bool = True
    redact_passport_numbers: bool = True
    redact_claim_numbers: bool = True
    redact_medical_record_numbers: bool = True
    custom_terms: tuple[str, ...] = ()


@dataclass
class RedactionResult:
    redacted_text: str
    replacement_counts: dict[str, int] = field(default_factory=dict)

    @property
    def total_replacements(self) -> int:
        return sum(self.replacement_counts.values())


EMAIL_RE = re.compile(r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b")
PHONE_RE = re.compile(
    r"\b(?:\+?1[-.\s]?)?(?:\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})\b"
)
SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
EIN_RE = re.compile(r"\b\d{2}-\d{7}\b")
DOB_RE = re.compile(r"\b(?:0?[1-9]|1[0-2])[/-](?:0?[1-9]|[12]\d|3[01])[/-](?:19|20)\d{2}\b")
PASSPORT_RE = re.compile(r"\b(?=[A-Z0-9]{6,9}\b)(?=[A-Z0-9]*[A-Z])(?=[A-Z0-9]*\d)[A-Z0-9]{6,9}\b")
CLAIM_RE = re.compile(r"\b(?:Claim|Case|Reference|Ref)\s*[:#-]?\s*[A-Z0-9\-]{6,}\b", re.IGNORECASE)
MRN_RE = re.compile(r"\b(?:MRN|Medical\s*Record\s*Number)\s*[:#-]?\s*[A-Z0-9\-]{4,}\b", re.IGNORECASE)
ADDRESS_RE = re.compile(
    r"\b\d{1,6}\s+[A-Za-z0-9.'\-\s]+\s(?:Street|St|Road|Rd|Avenue|Ave|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Way|Place|Pl)\b(?:,?\s+[A-Za-z.\s]+)?",
    re.IGNORECASE,
)
NAME_RE = re.compile(r"\b(?:[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b")


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _apply_pattern(text: str, pattern: re.Pattern[str], replacement: str) -> tuple[str, int]:
    return pattern.subn(replacement, text)


def _apply_custom_terms(text: str, custom_terms: Iterable[str]) -> tuple[str, int]:
    total = 0
    updated = text
    for term in custom_terms:
        clean = term.strip()
        if not clean:
            continue
        updated, count = re.subn(
            re.escape(clean), "[REDACTED_CUSTOM]", updated, flags=re.IGNORECASE
        )
        total += count
    return updated, total


def redact_text_with_report(text: str, options: RedactionOptions) -> RedactionResult:
    redacted = text
    counts: dict[str, int] = {}

    if options.redact_emails:
        redacted, count = _apply_pattern(redacted, EMAIL_RE, "[REDACTED_EMAIL]")
        counts["emails"] = count
    if options.redact_phones:
        redacted, count = _apply_pattern(redacted, PHONE_RE, "[REDACTED_PHONE]")
        counts["phones"] = count
    if options.redact_ssn_tax_ids:
        redacted, ssn_count = _apply_pattern(redacted, SSN_RE, "[REDACTED_SSN]")
        redacted, ein_count = _apply_pattern(redacted, EIN_RE, "[REDACTED_TAX_ID]")
        counts["ssn"] = ssn_count
        counts["tax_ids"] = ein_count
    if options.redact_dob:
        redacted, count = _apply_pattern(redacted, DOB_RE, "[REDACTED_DOB]")
        counts["dob"] = count
    if options.redact_passport_numbers:
        redacted, count = _apply_pattern(redacted, PASSPORT_RE, "[REDACTED_PASSPORT]")
        counts["passport_numbers"] = count
    if options.redact_claim_numbers:
        redacted, count = _apply_pattern(redacted, CLAIM_RE, "[REDACTED_CLAIM]")
        counts["claim_numbers"] = count
    if options.redact_medical_record_numbers:
        redacted, count = _apply_pattern(redacted, MRN_RE, "[REDACTED_MRN]")
        counts["medical_record_numbers"] = count
    if options.redact_addresses:
        redacted, count = _apply_pattern(redacted, ADDRESS_RE, "[REDACTED_ADDRESS]")
        counts["addresses"] = count
    if options.redact_names:
        redacted, count = _apply_pattern(redacted, NAME_RE, "[REDACTED_NAME]")
        counts["names"] = count

    redacted, custom_count = _apply_custom_terms(redacted, options.custom_terms)
    counts["custom_terms"] = custom_count

    return RedactionResult(redacted_text=redacted, replacement_counts=counts)


def redact_text(text: str, options: RedactionOptions) -> str:
    return redact_text_with_report(text, options).redacted_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def extract_text_from_pdf_with_ocr(file_bytes: bytes, enable_ocr_fallback: bool = False) -> tuple[str, bool]:
    text = extract_text_from_pdf(file_bytes)
    if text.strip() or not enable_ocr_fallback:
        return text, False

    try:
        import fitz  # pymupdf
        import pytesseract
    except Exception:
        return text, False

    pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
    all_text = []
    for page in pdf_doc:
        pix = page.get_pixmap(dpi=200)
        image_bytes = pix.tobytes("png")
        try:
            from PIL import Image
        except Exception:
            return text, False
        image = Image.open(io.BytesIO(image_bytes))
        ocr_text = pytesseract.image_to_string(image)
        all_text.append(ocr_text)

    ocr_combined = "\n\n".join(all_text)
    if ocr_combined.strip():
        return ocr_combined, True
    return text, False


def build_redacted_docx(text: str) -> bytes:
    from docx import Document

    document = Document()
    for line in text.splitlines() or [text]:
        document.add_paragraph(line)

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_redacted_pdf(text: str) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=LETTER)
    _, height = LETTER

    margin = 50
    y = height - margin
    for line in text.splitlines() or [text]:
        current_line = line
        while len(current_line) > 95:
            pdf.drawString(margin, y, current_line[:95])
            current_line = current_line[95:]
            y -= 14
            if y <= margin:
                pdf.showPage()
                y = height - margin
        pdf.drawString(margin, y, current_line)
        y -= 14
        if y <= margin:
            pdf.showPage()
            y = height - margin

    pdf.save()
    return stream.getvalue()


def redact_pdf_native(file_bytes: bytes, options: RedactionOptions) -> tuple[bytes | None, str | None]:
    """Native object-level redaction for PDFs using PyMuPDF if available.

    Returns (pdf_bytes, error_message). If unavailable/error, pdf_bytes is None.
    """
    try:
        import fitz
    except Exception:
        return None, "PyMuPDF is not installed; cannot run native PDF object redaction."

    doc = fitz.open(stream=file_bytes, filetype="pdf")

    regexes: list[tuple[re.Pattern[str], str]] = []
    if options.redact_emails:
        regexes.append((EMAIL_RE, "[REDACTED_EMAIL]"))
    if options.redact_phones:
        regexes.append((PHONE_RE, "[REDACTED_PHONE]"))
    if options.redact_ssn_tax_ids:
        regexes.extend([(SSN_RE, "[REDACTED_SSN]"), (EIN_RE, "[REDACTED_TAX_ID]")])
    if options.redact_dob:
        regexes.append((DOB_RE, "[REDACTED_DOB]"))
    if options.redact_claim_numbers:
        regexes.append((CLAIM_RE, "[REDACTED_CLAIM]"))
    if options.redact_medical_record_numbers:
        regexes.append((MRN_RE, "[REDACTED_MRN]"))

    for page in doc:
        page_text = page.get_text("text")
        terms: set[str] = set()
        for pattern, _ in regexes:
            for match in pattern.finditer(page_text):
                token = match.group(0).strip()
                if token:
                    terms.add(token)

        for term in options.custom_terms:
            clean = term.strip()
            if clean:
                terms.add(clean)

        for term in sorted(terms, key=len, reverse=True):
            for rect in page.search_for(term):
                page.add_redact_annot(rect, fill=(0, 0, 0))

        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    out = io.BytesIO()
    doc.save(out, garbage=4, deflate=True)
    return out.getvalue(), None
